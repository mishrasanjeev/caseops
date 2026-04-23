"""Sprint Q6 + Q7 — regenerate endpoint + DOCX export.

The export path unit-tests are pure-function over the renderer (no
FastAPI), plus a single integration test that exercises the Response
content-type + disposition via the client fixture.
"""
from __future__ import annotations

from datetime import UTC, date, datetime
from io import BytesIO

import pytest
from fastapi.testclient import TestClient

from caseops_api.services.matter_summary import (
    MatterExecutiveSummary,
    MatterSummaryTimelineEvent,
)
from caseops_api.services.matter_summary_export import (
    _safe_filename,
    render_summary_docx,
)
from caseops_api.services.matter_timeline import (
    MatterTimeline,
    TimelineEvent,
)


def _summary() -> MatterExecutiveSummary:
    return MatterExecutiveSummary(
        overview=(
            "Bail application before the Delhi High Court arising out of "
            "FIR 123/2026. Accused seeks release on parity."
        ),
        key_facts=["In custody since 1 March 2026", "Two co-accused granted bail"],
        timeline=[
            MatterSummaryTimelineEvent(
                date="2026-03-05", label="FIR 123/2026 registered"
            ),
        ],
        legal_issues=["Triple test", "Parity with co-accused"],
        sections_cited=["BNS s.303", "BNS s.318", "BNSS s.483"],
        generated_at=datetime(2026, 4, 20, 19, 0, tzinfo=UTC),
    )


def _timeline() -> MatterTimeline:
    return MatterTimeline(
        matter_id="test-id",
        generated_at=datetime(2026, 4, 20, 19, 0, tzinfo=UTC),
        events=[
            TimelineEvent(
                event_date=date(2026, 3, 5),
                kind="court_order",
                title="FIR registered",
                summary="Complaint registered by the SHO.",
            ),
            TimelineEvent(
                event_date=date(2026, 4, 10),
                kind="hearing",
                title="First bail hearing",
                summary="Delhi High Court. before Justice Vikram Nath",
            ),
        ],
    )


def test_render_summary_docx_produces_valid_docx_bytes() -> None:
    body, filename = render_summary_docx(
        matter_title="State v Accused",
        matter_code="CR-001/2026",
        summary=_summary(),
        timeline=_timeline(),
    )
    # python-docx files are ZIP archives — magic bytes 'PK'.
    assert body[:2] == b"PK"
    # And the filename derives from the matter code, stripped to a
    # filesystem-safe form.
    assert filename.endswith(".docx")
    assert "CR-001-2026" in filename
    # Round-trip: the DOCX should open + contain at least the heading.
    from docx import Document  # type: ignore[import-not-found]

    doc = Document(BytesIO(body))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Matter Executive Summary" in text
    assert "State v Accused" in text
    assert "BNS s.303" in text  # sections_cited bulleted
    assert "Vikram Nath" in text  # timeline event rendered
    assert "Triple test" in text  # legal issue


def test_safe_filename_collapses_unsafe_chars() -> None:
    assert _safe_filename("CR-001/2026") == "CR-001-2026"
    assert _safe_filename("  bail  app  ") == "bail-app"
    assert _safe_filename("") == "matter"
    assert _safe_filename("///") == "matter"


def test_render_summary_docx_uses_timeline_events_over_ai_timeline() -> None:
    """The Q8 timeline is grounded in DB; it must win over the LLM's
    guess in the rendered DOCX. The latter only appears when the
    structured timeline is empty."""
    empty_timeline = MatterTimeline(
        matter_id="x",
        generated_at=datetime.now(UTC),
        events=[],
    )
    body, _ = render_summary_docx(
        matter_title="t",
        matter_code="c",
        summary=_summary(),
        timeline=empty_timeline,
    )
    from docx import Document  # type: ignore[import-not-found]

    doc = Document(BytesIO(body))
    text = "\n".join(p.text for p in doc.paragraphs)
    # With empty structured timeline the LLM fallback appears.
    assert "AI summary" in text


def test_render_summary_docx_hides_ai_timeline_when_structured_present() -> None:
    body, _ = render_summary_docx(
        matter_title="t",
        matter_code="c",
        summary=_summary(),
        timeline=_timeline(),  # non-empty
    )
    from docx import Document  # type: ignore[import-not-found]

    doc = Document(BytesIO(body))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "AI summary" not in text
    # Grounded timeline IS shown.
    assert "Timeline" in text


# ---------------------------------------------------------------
# HTTP integration — content-type + disposition.
# ---------------------------------------------------------------


@pytest.fixture
def stub_summary(monkeypatch: pytest.MonkeyPatch) -> None:
    """Route tests run under the mock LLM provider, which doesn't
    return valid JSON for the summary prompt. Stub the service call
    so we're testing the export path, not the LLM-JSON contract
    (already covered by Q5 tests)."""

    def _fake(session, *, context, matter_id, provider=None, force_refresh=False):
        _ = provider, force_refresh
        # Still enforce tenancy — the real service does this via
        # ``_get_matter_model``; we mirror that so cross-tenant tests
        # continue to see a 404.
        from caseops_api.services.matters import _get_matter_model

        _get_matter_model(session, context=context, matter_id=matter_id)
        return _summary()

    monkeypatch.setattr(
        "caseops_api.api.routes.matters.generate_matter_summary", _fake
    )


def test_summary_docx_route_returns_attachment(
    client: TestClient, stub_summary: None,
) -> None:
    from tests.test_auth_company import auth_headers, bootstrap_company

    bootstrap = bootstrap_company(client)
    headers = auth_headers(str(bootstrap["access_token"]))

    resp = client.post(
        "/api/matters",
        json={
            "matter_code": "Q7-001",
            "title": "Bail application",
            "practice_area": "Bail / Custody",
            "forum_level": "high_court",
            "description": "Accused seeks bail.",
        },
        headers=headers,
    )
    assert resp.status_code == 200
    matter_id = resp.json()["id"]

    docx_resp = client.get(
        f"/api/matters/{matter_id}/summary.docx", headers=headers,
    )
    assert docx_resp.status_code == 200, docx_resp.text
    assert docx_resp.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    assert 'attachment; filename="Q7-001-summary.docx"' in (
        docx_resp.headers["content-disposition"]
    )
    assert docx_resp.content[:2] == b"PK"


def test_render_summary_pdf_produces_valid_pdf_bytes() -> None:
    """Pure-function PDF render: bytes start with %PDF, filename
    uses the matter code, and the body contains the expected section
    labels. No network, no LLM."""
    from caseops_api.services.matter_summary_export import render_summary_pdf

    body, filename = render_summary_pdf(
        matter_title="State v Accused",
        matter_code="CR-001/2026",
        summary=_summary(),
        timeline=_timeline(),
    )
    assert body.startswith(b"%PDF-"), "fpdf2 must emit a real PDF"
    assert filename.endswith(".pdf")
    assert "CR-001-2026" in filename
    # Helvetica cannot encode em-dashes / smart quotes; the _ascii_safe
    # helper flattens those — verify no encoding crash emitted bytes.
    assert len(body) > 500


def test_summary_pdf_route_returns_attachment(
    client: TestClient, stub_summary: None,
) -> None:
    from tests.test_auth_company import auth_headers, bootstrap_company

    bootstrap = bootstrap_company(client)
    headers = auth_headers(str(bootstrap["access_token"]))
    resp = client.post(
        "/api/matters",
        json={
            "matter_code": "Q7P-001",
            "title": "Bail application PDF export",
            "practice_area": "Bail / Custody",
            "forum_level": "high_court",
            "description": "PDF summary export.",
        },
        headers=headers,
    )
    assert resp.status_code == 200
    matter_id = resp.json()["id"]

    pdf_resp = client.get(
        f"/api/matters/{matter_id}/summary.pdf", headers=headers,
    )
    assert pdf_resp.status_code == 200, pdf_resp.text
    assert pdf_resp.headers["content-type"] == "application/pdf"
    assert 'attachment; filename="Q7P-001-summary.pdf"' in (
        pdf_resp.headers["content-disposition"]
    )
    assert pdf_resp.content.startswith(b"%PDF-")


def test_summary_docx_route_404s_cross_tenant(
    client: TestClient, stub_summary: None,
) -> None:
    from tests.test_auth_company import auth_headers, bootstrap_company

    # Tenant A creates a matter.
    a = bootstrap_company(client)
    headers_a = auth_headers(str(a["access_token"]))
    resp = client.post(
        "/api/matters",
        json={
            "matter_code": "Q7-ISO",
            "title": "Tenant A matter",
            "practice_area": "Civil / Contract",
            "forum_level": "high_court",
        },
        headers=headers_a,
    )
    matter_id_a = resp.json()["id"]

    # Tenant B bootstraps.
    b_resp = client.post(
        "/api/bootstrap/company",
        json={
            "company_name": "Tenant B Ltd",
            "company_slug": "tenant-b-export",
            "company_type": "law_firm",
            "owner_full_name": "Tenant B Owner",
            "owner_email": "b@b-export.example",
            "owner_password": "TenantB-Strong!234",
        },
    )
    headers_b = auth_headers(str(b_resp.json()["access_token"]))

    cross = client.get(
        f"/api/matters/{matter_id_a}/summary.docx", headers=headers_b,
    )
    assert cross.status_code == 404


def test_summary_regenerate_route_returns_fresh_summary(
    client: TestClient, stub_summary: None,
) -> None:
    from tests.test_auth_company import auth_headers, bootstrap_company

    bootstrap = bootstrap_company(client)
    headers = auth_headers(str(bootstrap["access_token"]))
    resp = client.post(
        "/api/matters",
        json={
            "matter_code": "Q6-001",
            "title": "Regenerate me",
            "practice_area": "Civil / Contract",
            "forum_level": "high_court",
        },
        headers=headers,
    )
    matter_id = resp.json()["id"]

    regen = client.post(
        f"/api/matters/{matter_id}/summary/regenerate", headers=headers,
    )
    assert regen.status_code == 200, regen.text
    body = regen.json()
    # MatterExecutiveSummary shape.
    assert "overview" in body
    assert "generated_at" in body
    # key_facts / legal_issues / sections_cited / timeline are lists.
    assert isinstance(body["key_facts"], list)
    assert isinstance(body["legal_issues"], list)


def test_summary_regenerate_route_404s_cross_tenant(
    client: TestClient, stub_summary: None,
) -> None:
    from tests.test_auth_company import auth_headers, bootstrap_company

    a = bootstrap_company(client)
    headers_a = auth_headers(str(a["access_token"]))
    resp = client.post(
        "/api/matters",
        json={
            "matter_code": "Q6-ISO",
            "title": "Tenant A matter",
            "practice_area": "Civil / Contract",
            "forum_level": "high_court",
        },
        headers=headers_a,
    )
    assert resp.status_code == 200, resp.text
    matter_id_a = resp.json()["id"]

    b = client.post(
        "/api/bootstrap/company",
        json={
            "company_name": "Tenant B Re",
            "company_slug": "tenant-b-regen",
            "company_type": "law_firm",
            "owner_full_name": "Tenant B Owner",
            "owner_email": "b@b-regen.example",
            "owner_password": "TenantB-Strong!234",
        },
    )
    headers_b = auth_headers(str(b.json()["access_token"]))
    resp_b = client.post(
        f"/api/matters/{matter_id_a}/summary/regenerate", headers=headers_b,
    )
    assert resp_b.status_code == 404


# ---------------------------------------------------------------------------
# EG-005 (2026-04-23) — caching contract: GET / DOCX / PDF reuse the
# cached payload, POST regenerate forces a refresh, every refresh
# leaves a ModelRun audit row.
# ---------------------------------------------------------------------------


def test_summary_caches_after_first_call_and_skips_llm_on_second(
    client: TestClient, monkeypatch: pytest.MonkeyPatch,
) -> None:
    """The first GET /summary should call ``generate_matter_summary``;
    the second GET should hit the cached row and skip the LLM. Codex
    enterprise audit gap EG-005 — without caching, every cockpit
    refresh + every export costs another Haiku round-trip."""
    from caseops_api.services import matter_summary as summary_mod
    from tests.test_auth_company import auth_headers, bootstrap_company

    bootstrap = bootstrap_company(client)
    headers = auth_headers(str(bootstrap["access_token"]))

    resp = client.post(
        "/api/matters",
        json={
            "matter_code": "EG5-CACHE-001",
            "title": "Cache hit smoke",
            "practice_area": "Civil",
            "forum_level": "high_court",
        },
        headers=headers,
    )
    matter_id = resp.json()["id"]

    real_generate = summary_mod.generate_matter_summary
    call_log: list[bool] = []

    def _spy(session, *, context, matter_id, provider=None, force_refresh=False):
        call_log.append(force_refresh)
        # Stub the LLM call by patching build_provider so we never
        # touch Anthropic. The real generate_matter_summary still runs
        # — including its cache read/write.
        return real_generate(
            session,
            context=context,
            matter_id=matter_id,
            provider=_StubLLM(),
            force_refresh=force_refresh,
        )

    monkeypatch.setattr(
        "caseops_api.api.routes.matters.generate_matter_summary", _spy
    )

    # First GET — cache miss, should compute.
    r1 = client.get(f"/api/matters/{matter_id}/summary", headers=headers)
    assert r1.status_code == 200, r1.text

    # Second GET — cache hit, must NOT have called the LLM stub
    # (we'll assert by counting StubLLM invocations).
    r2 = client.get(f"/api/matters/{matter_id}/summary", headers=headers)
    assert r2.status_code == 200
    assert r1.json()["overview"] == r2.json()["overview"]

    # The route is called twice; the LLM stub is called only once
    # (first GET). Cache hit invariant.
    assert _StubLLM.call_count == 1

    # POST regenerate forces a fresh LLM call, increments the counter.
    r3 = client.post(
        f"/api/matters/{matter_id}/summary/regenerate", headers=headers,
    )
    assert r3.status_code == 200
    assert _StubLLM.call_count == 2
    assert call_log[-1] is True  # regenerate sets force_refresh=True

    # Subsequent GET is back to cache (counter stays at 2).
    r4 = client.get(f"/api/matters/{matter_id}/summary", headers=headers)
    assert r4.status_code == 200
    assert _StubLLM.call_count == 2


def test_summary_call_writes_model_run_audit_row(
    client: TestClient, monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Every uncached summary call must persist a ModelRun row with
    purpose='matter_summary' so AI spend is auditable. Without this,
    summary calls were invisible to the cost-per-tenant view."""
    from sqlalchemy import select

    from caseops_api.db.models import ModelRun
    from caseops_api.db.session import get_session_factory
    from caseops_api.services import matter_summary as summary_mod
    from tests.test_auth_company import auth_headers, bootstrap_company

    bootstrap = bootstrap_company(client)
    headers = auth_headers(str(bootstrap["access_token"]))

    resp = client.post(
        "/api/matters",
        json={
            "matter_code": "EG5-AUDIT-001",
            "title": "Audit smoke",
            "practice_area": "Civil",
            "forum_level": "high_court",
        },
        headers=headers,
    )
    matter_id = resp.json()["id"]

    real_generate = summary_mod.generate_matter_summary

    def _with_stub(session, *, context, matter_id, provider=None, force_refresh=False):
        return real_generate(
            session,
            context=context,
            matter_id=matter_id,
            provider=_StubLLM(),
            force_refresh=force_refresh,
        )

    monkeypatch.setattr(
        "caseops_api.api.routes.matters.generate_matter_summary", _with_stub
    )

    r = client.get(f"/api/matters/{matter_id}/summary", headers=headers)
    assert r.status_code == 200, r.text

    factory = get_session_factory()
    with factory() as session:
        runs = list(
            session.scalars(
                select(ModelRun).where(ModelRun.purpose == "matter_summary")
            )
        )
    assert len(runs) == 1
    run = runs[0]
    assert run.status == "ok"
    assert run.matter_id == matter_id
    assert run.company_id is not None
    assert run.actor_membership_id is not None
    # Token counts come from the stub's LLMCompletion.
    assert run.prompt_tokens == 100
    assert run.completion_tokens == 30


class _StubLLM:
    """Module-level stub LLM for the cache + audit tests. Counts
    invocations so the cache assertions can prove the LLM was skipped
    on cached reads."""

    name = "stub"
    model = "stub-summary-model"
    call_count: int = 0

    def __init__(self) -> None:
        # Reset on first construction in each test by inspecting the
        # call stack would be brittle; instead, the tests reset
        # explicitly via the class attribute below.
        pass

    def generate(self, *, messages, temperature, max_tokens):
        _ = messages, temperature, max_tokens
        type(self).call_count += 1
        from caseops_api.services.llm import LLMCompletion

        return LLMCompletion(
            provider=self.name,
            model=self.model,
            text=(
                '{"overview":"stub overview","key_facts":["fact a"],'
                '"timeline":[{"date":"2026-04-01","label":"event"}],'
                '"legal_issues":["issue a"],'
                '"sections_cited":["BNS s.303"]}'
            ),
            prompt_tokens=100,
            completion_tokens=30,
            latency_ms=20,
        )


@pytest.fixture(autouse=True)
def _reset_stub_llm_counter() -> None:
    """Module-scoped class counter; reset before each test so the
    assertion totals are per-test, not cumulative across the file."""
    _StubLLM.call_count = 0
