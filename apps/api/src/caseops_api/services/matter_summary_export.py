"""Sprint Q7 — export the matter executive summary as DOCX.

Takes an existing ``MatterExecutiveSummary`` (produced by
``services.matter_summary``) plus the matter's ``MatterTimeline``
(Sprint Q8) and renders both into a ready-to-download DOCX.

PDF export is a planned follow-up once we add ``fpdf2`` to the
dependency set — shipping DOCX now avoids a new native-dep ask and
covers ~80 % of the filing-ready use-case (lawyers routinely circulate
DOCX for markup).

The renderer is a pure function over the input dataclasses so it is
easy to unit-test without a live LLM call.
"""
from __future__ import annotations

import re
from io import BytesIO

from caseops_api.services.matter_summary import MatterExecutiveSummary
from caseops_api.services.matter_timeline import MatterTimeline


def render_summary_docx(
    *,
    matter_title: str,
    matter_code: str,
    summary: MatterExecutiveSummary,
    timeline: MatterTimeline,
) -> tuple[bytes, str]:
    """Return ``(bytes, filename)`` for the summary DOCX.

    Filename is derived from ``matter_code`` so it lands in the
    caller's downloads folder with an obvious association to the
    matter. Whitespace / slashes in the code are collapsed to hyphens.
    """
    from docx import Document  # type: ignore[import-not-found]
    from docx.shared import Pt

    doc = Document()

    title = doc.add_heading("Matter Executive Summary", level=1)
    for run in title.runs:
        run.font.size = Pt(20)

    meta = doc.add_paragraph()
    meta.add_run(f"{matter_title} ({matter_code})").bold = True
    meta.add_run("  ·  ")
    meta.add_run(
        f"Generated {summary.generated_at.strftime('%d %b %Y, %H:%M')}"
    ).italic = True

    doc.add_paragraph()  # spacer

    if summary.overview:
        doc.add_heading("Overview", level=2)
        doc.add_paragraph(summary.overview)

    if summary.key_facts:
        doc.add_heading("Key facts", level=2)
        for fact in summary.key_facts:
            doc.add_paragraph(fact, style="List Bullet")

    if summary.legal_issues:
        doc.add_heading("Legal issues", level=2)
        for issue in summary.legal_issues:
            doc.add_paragraph(issue, style="List Bullet")

    if summary.sections_cited:
        doc.add_heading("Statutes and sections cited", level=2)
        for sec in summary.sections_cited:
            doc.add_paragraph(sec, style="List Bullet")

    # Prefer the merged Q8 timeline over the summary's own guess —
    # Q8 is grounded in hearings / deadlines / court orders, not an
    # LLM hallucination.
    if timeline.events:
        doc.add_heading("Timeline", level=2)
        for event in timeline.events:
            p = doc.add_paragraph(style="List Number")
            p.add_run(f"{event.event_date.isoformat()} — ").bold = True
            p.add_run(f"{event.title}. ")
            p.add_run(event.summary)
    elif summary.timeline:
        # Fallback: LLM-generated timeline if there's nothing structured.
        doc.add_heading("Timeline (AI summary)", level=2)
        for event in summary.timeline:
            p = doc.add_paragraph(style="List Number")
            prefix = event.date or "—"
            p.add_run(f"{prefix} — ").bold = True
            p.add_run(event.label)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(
        "AI-assisted summary. Every cited statute traces to documents "
        "attached to this matter. Review before filing or circulating."
    ).italic = True

    buf = BytesIO()
    doc.save(buf)
    filename = _safe_filename(matter_code) + "-summary.docx"
    return buf.getvalue(), filename


def render_summary_pdf(
    *,
    matter_title: str,
    matter_code: str,
    summary: MatterExecutiveSummary,
    timeline: MatterTimeline,
) -> tuple[bytes, str]:
    """Sprint Q7 PDF slice — same shape as ``render_summary_docx``
    but emits a PDF using ``fpdf2`` (pure Python MIT, no native deps).

    Prefers the grounded Q8 timeline over the LLM-guessed one in the
    summary payload, matching the DOCX path.
    """
    from fpdf import FPDF  # type: ignore[import-not-found]

    pdf = FPDF(format="A4", unit="mm")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", size=18, style="B")
    pdf.cell(0, 12, _ascii_safe("Matter Executive Summary"), new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("Helvetica", size=11, style="B")
    pdf.cell(
        0, 7,
        _ascii_safe(f"{matter_title} ({matter_code})"),
        new_x="LMARGIN", new_y="NEXT",
    )
    pdf.set_font("Helvetica", size=9, style="I")
    pdf.cell(
        0, 5,
        _ascii_safe(f"Generated {summary.generated_at.strftime('%d %b %Y, %H:%M')}"),
        new_x="LMARGIN", new_y="NEXT",
    )
    pdf.ln(3)

    if summary.overview:
        _pdf_section(pdf, "Overview", [summary.overview])

    if summary.key_facts:
        _pdf_section(pdf, "Key facts", summary.key_facts, bullet=True)

    if summary.legal_issues:
        _pdf_section(pdf, "Legal issues", summary.legal_issues, bullet=True)

    if summary.sections_cited:
        _pdf_section(pdf, "Statutes and sections cited", summary.sections_cited, bullet=True)

    if timeline.events:
        _pdf_heading(pdf, "Timeline")
        for event in timeline.events:
            pdf.set_font("Helvetica", size=10, style="B")
            pdf.multi_cell(
                0, 5,
                _ascii_safe(f"{event.event_date.isoformat()} — {event.title}"),
                new_x="LMARGIN", new_y="NEXT",
            )
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 5, _ascii_safe(event.summary), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
    elif summary.timeline:
        _pdf_heading(pdf, "Timeline (AI summary)")
        for event in summary.timeline:
            pdf.set_font("Helvetica", size=10, style="B")
            prefix = event.date or "—"
            pdf.multi_cell(
                0, 5,
                _ascii_safe(f"{prefix} — {event.label}"),
                new_x="LMARGIN", new_y="NEXT",
            )
            pdf.ln(1)

    pdf.ln(3)
    pdf.set_font("Helvetica", size=9, style="I")
    pdf.multi_cell(
        0, 5,
        _ascii_safe(
            "AI-assisted summary. Every cited statute traces to documents "
            "attached to this matter. Review before filing or circulating."
        ),
        new_x="LMARGIN", new_y="NEXT",
    )

    body = bytes(pdf.output())
    filename = _safe_filename(matter_code) + "-summary.pdf"
    return body, filename


def _pdf_heading(pdf, text: str) -> None:
    pdf.set_font("Helvetica", size=13, style="B")
    pdf.ln(2)
    pdf.cell(0, 7, _ascii_safe(text), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", size=10)


def _pdf_section(pdf, heading: str, items: list[str], *, bullet: bool = False) -> None:
    _pdf_heading(pdf, heading)
    for item in items:
        if bullet:
            pdf.multi_cell(
                0, 5,
                _ascii_safe(f"• {item}"),
                new_x="LMARGIN", new_y="NEXT",
            )
        else:
            pdf.multi_cell(0, 5, _ascii_safe(item), new_x="LMARGIN", new_y="NEXT")


def _ascii_safe(text: str) -> str:
    """fpdf2's built-in Helvetica is a WinAnsi font — non-Latin code
    points blow up the PDF writer. For Q7's first PDF slice we flatten
    to an ASCII-biased fallback; Hindi / Tamil / etc bodies still ship
    through the DOCX path unchanged, and a future pass can add a
    TrueType font subset for the PDF path."""
    if not text:
        return ""
    # Replace typographic dashes / ellipses / quotes with ASCII
    # equivalents, then best-effort drop anything Helvetica-1252
    # can't encode.
    table = {
        "\u2013": "-",  # en dash
        "\u2014": "--",  # em dash
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2026": "...",
        "\u00a0": " ",  # non-breaking space
        "•": "*",
    }
    out = text.translate({ord(k): v for k, v in table.items()})
    try:
        out.encode("latin-1")
    except UnicodeEncodeError:
        out = out.encode("ascii", "replace").decode("ascii")
    return out


def _safe_filename(matter_code: str) -> str:
    """Collapse whitespace / slashes to hyphens so the download is
    safe on every OS."""
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", (matter_code or "matter").strip())
    return cleaned.strip("-") or "matter"


__all__ = ["render_summary_docx", "render_summary_pdf"]
