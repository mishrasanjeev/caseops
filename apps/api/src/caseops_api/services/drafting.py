"""Drafting studio pipeline (PRD §9.5, §10.3).

Owns the full state machine for a legal draft:

    draft (empty)  ── generate ──>  draft (v1)
                                      │
                                      ▼
                                   in_review  ── request_changes ──> changes_requested
                                      │                                      │
                                      │                         regenerate ──┘
                                      ▼
                                   approved  ── finalize ──> finalized (terminal)

Rules that ship with v1:

- Every new or regenerated version must emit at least one citation that
  survives the citation verifier (``services/citations.verify_citations``).
  ``approve`` fails closed when the current version has zero verified
  citations — no silent "approved without sources".
- Only a draft in ``in_review`` may transition to ``approved``; only
  ``approved`` can finalize. All transitions are recorded as a
  ``DraftReview`` row so the audit trail shows who moved it.
- ``finalized`` is terminal. Regeneration and submission are refused.
- ``review_required`` stays ``True`` on every status except ``approved``
  and ``finalized``. The UI uses this to lock external-share actions.

The LLM call is the same provider abstraction the recommendation and
hearing-pack services use. ``MockProvider`` emits a deterministic
drafting JSON so the full pipeline is exercisable offline.
"""
from __future__ import annotations

import hashlib
import io
import json
import logging
import re
from datetime import UTC, datetime
from typing import Literal

from fastapi import HTTPException, status
from pydantic import BaseModel, Field, ValidationError
from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload

from caseops_api.db.models import (
    AuthorityDocument,
    Draft,
    DraftReview,
    DraftReviewAction,
    DraftStatus,
    DraftType,
    DraftVersion,
    Matter,
    ModelRun,
)
from caseops_api.services.audit import record_from_context
from caseops_api.services.authorities import search_authority_catalog
from caseops_api.services.citations import (
    Claim,
    SourceDoc,
    VerificationReport,
    verify_citations,
)
from caseops_api.services.draft_validators import (
    DraftFinding,
    run_validators,
)
from caseops_api.services.identity import SessionContext
from caseops_api.services.llm import (
    PURPOSE_DRAFTING,
    AnthropicProvider,
    LLMCallContext,
    LLMCompletion,
    LLMMessage,
    LLMProvider,
    LLMProviderError,
    LLMQuotaExhaustedError,
    OpenAIProvider,
    build_provider,
    generate_structured,
    max_tokens_for_purpose,
)
from caseops_api.services.matter_access import assert_access

logger = logging.getLogger(__name__)


# Same Haiku fallback pattern as services.recommendations /
# services.matter_summary. Sonnet 4.6 sporadically returns malformed
# JSON on long structured outputs (observed 2026-04-20 on the
# recommendations endpoint, also flagged by end users as BUG-001 /
# BUG-002 — drafts "not created, no error toast, indefinite hang").
# Haiku is materially more reliable on JSON shape; we retry once
# before raising 422 so the user always gets either a draft or a
# clear error, never a silent hang.
_HAIKU_FALLBACK_MODEL = "claude-haiku-4-5-20251001"


def _haiku_fallback_provider() -> LLMProvider | None:
    from caseops_api.core.settings import get_settings

    settings = get_settings()
    if (settings.llm_provider or "").lower() != "anthropic":
        return None
    return AnthropicProvider(
        model=_HAIKU_FALLBACK_MODEL,
        api_key=settings.llm_api_key,
        prompt_cache=bool(getattr(settings, "llm_prompt_cache_enabled", True)),
    )


def _openai_fallback_provider() -> LLMProvider | None:
    """Cross-provider hard cutover for Anthropic 402 ("credit balance
    is too low") events. Returns None when no OpenAI key is configured
    so local / dev / test runs continue to hit the existing 422 path
    without surprise outbound calls."""
    from caseops_api.core.settings import get_settings

    settings = get_settings()
    if not getattr(settings, "openai_api_key", None):
        return None
    return OpenAIProvider(
        model=getattr(settings, "openai_fallback_model", "gpt-5.1"),
        api_key=settings.openai_api_key,
    )


def _generate_draft_via_openai(invoke, root_exc: Exception):
    """Last-chance OpenAI cutover for the drafting pipeline.

    Either succeeds (returns ``(response, completion)``) or raises a 422
    HTTPException with an actionable detail. ``root_exc`` is the original
    failure we're cutting over from, included in the user-visible message
    so support can correlate from a single screenshot."""
    openai = _openai_fallback_provider()
    if openai is None:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "Could not generate a draft: the primary model failed "
                f"({type(root_exc).__name__}) and no OpenAI fallback is "
                "configured. Please retry in a minute, or contact support."
            ),
        ) from root_exc
    try:
        return invoke(openai)
    except (LLMProviderError, ValidationError) as oa_exc:
        logger.warning("Draft OpenAI fallback also failed: %s", oa_exc)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "Could not generate a draft: both Anthropic "
                f"({type(root_exc).__name__}) and the OpenAI fallback "
                f"({type(oa_exc).__name__}) failed. Please retry in a "
                "minute, or contact support if this persists."
            ),
        ) from oa_exc


PURPOSE = "draft"


class _LLMDraftResponse(BaseModel):
    body: str = Field(min_length=20, max_length=20000)
    citations: list[str] = Field(default_factory=list, max_length=30)
    summary: str | None = Field(default=None, max_length=1200)


def _load_matter(session: Session, context: SessionContext, matter_id: str) -> Matter:
    matter = session.scalar(
        select(Matter).where(
            Matter.id == matter_id, Matter.company_id == context.company.id
        )
    )
    if matter is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Matter not found."
        )
    assert_access(session, context=context, matter=matter)
    return matter


def _load_draft(
    session: Session, matter: Matter, draft_id: str, *, include_children: bool = True
) -> Draft:
    query = select(Draft).where(
        Draft.id == draft_id, Draft.matter_id == matter.id
    )
    if include_children:
        query = query.options(
            selectinload(Draft.versions),
            selectinload(Draft.reviews),
        )
    draft = session.scalar(query)
    if draft is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND, detail="Draft not found."
        )
    return draft


def create_draft(
    session: Session,
    *,
    context: SessionContext,
    matter_id: str,
    title: str,
    draft_type: str = DraftType.BRIEF,
    template_type: str | None = None,
    facts: dict | None = None,
) -> Draft:
    matter = _load_matter(session, context, matter_id)
    facts_json = None
    if facts:
        # Defensive cap: the facts dict is user-entered per-field text
        # (addresses, names, amounts, …). 64 KiB of JSON is ample for
        # every template we ship today and bounds worst-case row size.
        serialised = json.dumps(facts, ensure_ascii=False)
        if len(serialised) > 64 * 1024:
            raise HTTPException(
                status_code=status.HTTP_413_CONTENT_TOO_LARGE,
                detail="Facts payload exceeds 64 KiB.",
            )
        facts_json = serialised
    draft = Draft(
        matter_id=matter.id,
        created_by_membership_id=context.membership.id,
        title=title.strip(),
        draft_type=draft_type,
        template_type=template_type,
        status=DraftStatus.DRAFT,
        review_required=True,
        facts_json=facts_json,
    )
    session.add(draft)
    session.flush()
    record_from_context(
        session,
        context,
        action="draft.created",
        target_type="draft",
        target_id=draft.id,
        matter_id=matter.id,
        metadata={
            "title": draft.title,
            "draft_type": draft.draft_type,
            "template_type": draft.template_type,
            "facts_keys": sorted(facts.keys()) if facts else [],
        },
    )
    session.commit()
    session.refresh(draft)
    return draft


def list_drafts(
    session: Session,
    *,
    context: SessionContext,
    matter_id: str,
) -> list[Draft]:
    matter = _load_matter(session, context, matter_id)
    return list(
        session.scalars(
            select(Draft)
            .where(Draft.matter_id == matter.id)
            .options(selectinload(Draft.versions), selectinload(Draft.reviews))
            .order_by(Draft.updated_at.desc(), Draft.id.desc())
        )
    )


def get_draft(
    session: Session,
    *,
    context: SessionContext,
    matter_id: str,
    draft_id: str,
) -> Draft:
    matter = _load_matter(session, context, matter_id)
    return _load_draft(session, matter, draft_id)


_STATUTE_GUIDANCE = (
    "Indian statute disambiguation — apply strictly:\n"
    "- BNS (Bharatiya Nyaya Sanhita, 2023) is the substantive criminal code "
    "(successor to the IPC). Its sections define offences.\n"
    "- BNSS (Bharatiya Nagarik Suraksha Sanhita, 2023) is the procedural "
    "criminal code (successor to the CrPC). Its sections govern procedure — "
    "including bail, arrest, remand, investigation, and trial. "
    "In particular: bail after arrest is BNSS s.483 (~ CrPC s.439); "
    "anticipatory bail is BNSS s.482 (~ CrPC s.438); default bail is "
    "BNSS s.187 (~ CrPC s.167(2)).\n"
    "- BSA (Bharatiya Sakshya Adhiniyam, 2023) is the evidence code "
    "(successor to the Indian Evidence Act).\n"
    "Do NOT cite BNS for procedural provisions. Do NOT cite BNSS for "
    "substantive offence definitions. If the focus note references a "
    "section number, verify the statute against this list before writing "
    "the section into the body."
)


def _build_messages(
    matter: Matter,
    draft: Draft,
    retrieved: list[AuthorityDocument],
    focus_note: str | None,
) -> list[LLMMessage]:
    system = (
        "You are drafting a legal document for an Indian litigation "
        "matter.\n\n"
        "Output strictly valid JSON shaped as "
        "{\"body\": string, \"citations\": string[], \"summary\": string?}. "
        "No prose, no markdown fences.\n\n"
        "ABSOLUTE RULES — VIOLATING ANY OF THESE FAILS THE DRAFT:\n"
        "1. Do NOT invent facts. The only permissible facts are those "
        "present in the MATTER RECORD block below and the FOCUS NOTE. "
        "For any fact not explicitly supplied — FIR number, arrest date, "
        "witness identity, address, age, family composition, amounts, "
        "dates — write a square-bracket placeholder such as `[____]`, "
        "`[date of arrest]`, `[FIR number]`, `[address]`. Placeholders "
        "are expected and preferred; invented specifics are a defect.\n"
        "2. Do NOT invent authorities. The ONLY citable authorities are "
        "those in the AUTHORITIES block below. Cite each one inline in "
        "the body the first time you rely on it, using the exact "
        "identifier in square brackets — for example [2023:DHC:8921] or "
        "[(2022) 10 SCC 51]. Every substantive legal proposition must be "
        "anchored to at least one such citation. If no retrieved "
        "authority supports a proposition, drop the proposition or flag "
        "it in the summary; do not paper over it.\n"
        "3. Get the statute right. Read the STATUTE GUIDANCE block "
        "carefully; Bharatiya Nyaya Sanhita (BNS) and Bharatiya Nagarik "
        "Suraksha Sanhita (BNSS) are distinct. Confusing them is a "
        "disqualifying error.\n"
        "4. Keep register formal and paragraphs short. Preserve "
        "Indian-English conventions. The body must be a complete "
        "document (cause-title, parties, facts, grounds, prayer, "
        "verification as applicable to the draft type) — not an outline."
    )

    parts: list[str] = [_STATUTE_GUIDANCE, ""]
    parts.append("=== MATTER RECORD (only source of facts) ===")
    parts.append(f"Matter: {matter.title} ({matter.matter_code})")
    parts.append(f"Practice area: {matter.practice_area}")
    parts.append(f"Forum: {matter.forum_level}")
    if matter.court_name:
        parts.append(f"Court: {matter.court_name}")
    if matter.judge_name:
        parts.append(f"Judge: {matter.judge_name}")
    if matter.client_name:
        parts.append(f"Client: {matter.client_name}")
    if matter.opposing_party:
        parts.append(f"Opposing party: {matter.opposing_party}")
    if matter.description:
        parts.append(f"Background: {matter.description}")
    parts.append(
        "(Any fact not listed above must appear in the draft as a "
        "bracketed placeholder — never as an invented specific.)"
    )
    parts.append("")
    parts.append(f"Draft title: {draft.title}")
    parts.append(f"Draft type: {draft.draft_type}")
    if draft.template_type:
        parts.append(f"Template: {draft.template_type}")
    if draft.facts_json:
        try:
            facts = json.loads(draft.facts_json)
        except json.JSONDecodeError:
            facts = None
        if isinstance(facts, dict) and facts:
            parts.append("")
            parts.append(
                "=== STEPPER FACTS (authoritative — use as-is) ==="
            )
            for key, value in facts.items():
                if value is None or value == "":
                    continue
                parts.append(f"- {key}: {value}")
            parts.append(
                "(These facts came from the drafting stepper; treat them "
                "as the lawyer's verified input. Do NOT override with "
                "bracketed placeholders for any field listed above.)"
            )
    if focus_note:
        parts.append(f"Focus: {focus_note}")
    parts.append("")

    citable = [doc for doc in retrieved if (doc.neutral_citation or doc.case_reference)]
    uncitable = [doc for doc in retrieved if not (doc.neutral_citation or doc.case_reference)]

    if citable:
        parts.append("=== AUTHORITIES (cite these — and ONLY these — inline) ===")
        for doc in citable:
            ident = doc.neutral_citation or doc.case_reference
            parts.append(f"- CITATION: {ident}")
            if doc.title:
                parts.append(f"  TITLE: {doc.title[:200]}")
            if doc.summary:
                excerpt = doc.summary.strip().splitlines()[0][:300]
                parts.append(f"  EXCERPT: {excerpt}")
        if uncitable:
            parts.append(
                f"(Additional {len(uncitable)} retrieved document(s) lack a "
                "reportable citation and are excluded from the citable set.)"
            )
    elif retrieved:
        parts.append(
            "=== AUTHORITIES ===\n"
            "The retrieval hit relevant documents but none carry a "
            "reportable citation (neutral citation or case reference). "
            "Do not cite by UUID or internal id. Produce the draft with "
            "bracketed `[citation needed]` anchors where authority "
            "should appear, and flag this gap in the summary."
        )
    else:
        parts.append(
            "=== AUTHORITIES ===\n"
            "No authorities retrieved. Produce a draft that flags "
            "`missing authorities` in the summary rather than inventing "
            "sources. Use `[citation needed]` anchors where authority "
            "should appear."
        )

    parts.append("")
    parts.append(
        "Respond with json. The citations array must list only identifiers "
        "you actually cited inline in the body."
    )

    return [
        LLMMessage(role="system", content=system),
        LLMMessage(role="user", content="\n".join(parts)),
    ]


def _prompt_hash(messages: list[LLMMessage]) -> str:
    digest = hashlib.sha256()
    for m in messages:
        digest.update(m.role.encode("utf-8"))
        digest.update(b"\x1f")
        digest.update(m.content.encode("utf-8"))
        digest.update(b"\x1e")
    return digest.hexdigest()


def _write_model_run(
    session: Session,
    *,
    context: SessionContext,
    matter_id: str,
    completion: LLMCompletion,
    prompt_hash: str,
    status_label: str = "ok",
    error: str | None = None,
) -> ModelRun:
    run = ModelRun(
        company_id=context.company.id,
        matter_id=matter_id,
        actor_membership_id=context.membership.id,
        purpose=PURPOSE,
        provider=completion.provider,
        model=completion.model,
        prompt_hash=prompt_hash,
        prompt_tokens=completion.prompt_tokens,
        completion_tokens=completion.completion_tokens,
        latency_ms=completion.latency_ms,
        status=status_label,
        error=error,
    )
    session.add(run)
    session.flush()
    return run


_SUMMARY_MAX_LEN = 1200


# Seed query packs keyed by a normalised practice signal. Each pack is
# merged into retrieval (multi-query → union + dedup) so the draft has a
# reasonable shot at inline-citing canonical precedents even when the
# matter title is terse. Keep packs small (3–5 queries) to stay under
# retrieval-cost limits.
_RETRIEVAL_PACKS: dict[str, list[str]] = {
    "bail": [
        "triple test for bail flight risk tampering evidence repetition offence",
        "parity co-accused bail granted identical footing",
        "prolonged custody undertrial pretrial detention bail default",
        "bail is the rule jail is the exception BNSS procedure",
    ],
    "anticipatory_bail": [
        "anticipatory bail custodial interrogation BNSS Section 482",
        "pre-arrest bail economic offence liberty",
    ],
    "quashing": [
        "quashing FIR Section 482 CrPC inherent powers BNSS 528",
        "abuse of process of court quashing proceedings",
    ],
}

_BAIL_HINTS = re.compile(
    r"\b(bail|BNSS\s*4[78]\d|CrPC\s*(438|439|167)|custody|undertrial)\b", re.I
)
_ANTICIPATORY_HINTS = re.compile(r"\banticipatory\s*bail\b|pre-arrest\s*bail", re.I)
_QUASHING_HINTS = re.compile(r"\bquash(ing)?\b|\bFIR\s*quash", re.I)


def _retrieval_queries(matter: Matter, focus_note: str | None) -> list[str]:
    """Build a prioritised list of queries to run against the corpus.

    The first entry is the "direct" matter query; subsequent entries
    are seed queries drawn from practice-area-specific packs. The
    drafting pipeline unions the retrieval results and de-duplicates by
    document id, so an empty pack is harmless.
    """
    base_parts: list[str] = [matter.title or ""]
    if matter.description:
        base_parts.append(matter.description)
    if focus_note:
        base_parts.append(focus_note)
    base = " ".join(p.strip() for p in base_parts if p and p.strip()).strip()
    queries: list[str] = [base] if base else []

    probe = " ".join([base, matter.practice_area or "", (matter.description or "")])
    added: set[str] = set()

    def _add(pack: str) -> None:
        if pack in added:
            return
        queries.extend(_RETRIEVAL_PACKS.get(pack, []))
        added.add(pack)

    if _ANTICIPATORY_HINTS.search(probe):
        _add("anticipatory_bail")
    if _BAIL_HINTS.search(probe):
        _add("bail")
    if _QUASHING_HINTS.search(probe):
        _add("quashing")

    return queries


def _retrieve_for_draft(
    session: Session,
    matter: Matter,
    focus_note: str | None,
    per_query_limit: int = 5,
    overfetch_limit: int = 16,
    final_top_k: int = 5,
) -> list[AuthorityDocument]:
    """Run multi-query retrieval, dedup, then rerank down to ``final_top_k``.

    Pipeline: first-stage hybrid retrieval over-fetches `overfetch_limit`
    candidates across the seed query packs; the reranker (mock by
    default, LLM-judge when ``CASEOPS_RERANK_ENABLED=true``) scores
    them against the matter context and returns the top K.
    """
    from caseops_api.services.reranker import (
        RerankerCandidate,
        build_reranker,
    )

    queries = _retrieval_queries(matter, focus_note)
    base_query = queries[0] if queries else ""

    seen: dict[str, AuthorityDocument] = {}
    for query in queries:
        if not query or len(seen) >= overfetch_limit:
            continue
        hits = search_authority_catalog(session, query=query, limit=per_query_limit)
        for hit in hits:
            if hit.authority_document_id in seen:
                continue
            doc = session.get(AuthorityDocument, hit.authority_document_id)
            if doc is not None:
                seen[hit.authority_document_id] = doc
                if len(seen) >= overfetch_limit:
                    break

    docs = list(seen.values())
    if len(docs) <= final_top_k:
        return docs

    reranker = build_reranker()
    cands = [
        RerankerCandidate(
            identifier=doc.id,
            title=doc.title or "",
            text=(doc.summary or "")[:500],
        )
        for doc in docs
    ]
    ranked = reranker.rerank(base_query or "legal research", cands, top_k=final_top_k)
    # Preserve reranker-determined order.
    by_id = {doc.id: doc for doc in docs}
    return [by_id[c.identifier] for c in ranked if c.identifier in by_id][:final_top_k]


def _augment_summary_with_findings(
    base: str | None, findings: list[DraftFinding]
) -> str | None:
    if not findings:
        return base
    lines = [f"[{f.severity.upper()}] {f.code}: {f.message}" for f in findings]
    suffix = "Review findings:\n" + "\n".join(lines)
    if not base:
        combined = suffix
    else:
        combined = f"{base}\n\n{suffix}"
    if len(combined) > _SUMMARY_MAX_LEN:
        combined = combined[: _SUMMARY_MAX_LEN - 3].rstrip() + "..."
    return combined


def _verify_version_citations(
    session: Session,
    citations: list[str],
) -> tuple[list[str], int]:
    """Verify that each citation is an authority we actually hold.
    Returns (surviving_citations, verified_count)."""
    unique = list(dict.fromkeys(c.strip() for c in citations if c and c.strip()))
    if not unique:
        return [], 0
    docs = list(
        session.scalars(
            select(AuthorityDocument).where(
                (AuthorityDocument.neutral_citation.in_(unique))
                | (AuthorityDocument.case_reference.in_(unique))
                | (AuthorityDocument.id.in_(unique))
            )
        )
    )
    sources: list[SourceDoc] = []
    known: set[str] = set()
    for doc in docs:
        identifier = doc.neutral_citation or doc.case_reference or doc.id
        aliases = tuple(
            filter(None, {doc.id, doc.neutral_citation, doc.case_reference})
        )
        sources.append(
            SourceDoc(identifier=identifier, aliases=aliases, text=doc.summary or "")
        )
        known.add(identifier)
        known.add(doc.id)
        if doc.neutral_citation:
            known.add(doc.neutral_citation)
        if doc.case_reference:
            known.add(doc.case_reference)
    claims = [Claim(citation=c) for c in unique]
    report: VerificationReport = verify_citations(claims, sources)
    surviving = [c for c in unique if c in known]
    return surviving, report.verified_count


def generate_draft_version(
    session: Session,
    *,
    context: SessionContext,
    matter_id: str,
    draft_id: str,
    focus_note: str | None = None,
    template_key: str | None = None,
    provider: LLMProvider | None = None,
) -> Draft:
    del template_key  # reserved for future template selection
    matter = _load_matter(session, context, matter_id)
    draft = _load_draft(session, matter, draft_id)

    if draft.status == DraftStatus.FINALIZED:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Finalized drafts cannot be regenerated.",
        )

    retrieved_docs = _retrieve_for_draft(session, matter, focus_note)

    messages = _build_messages(matter, draft, retrieved_docs, focus_note)
    prompt_hash = _prompt_hash(messages)
    # Drafting routes to the per-purpose drafting model (Opus-class
    # when configured); metadata extraction and recommendations pick
    # their own tier via build_provider(purpose=...).
    llm = provider or build_provider(purpose=PURPOSE_DRAFTING)
    llm_context = LLMCallContext(
        tenant_id=context.company.id, matter_id=matter.id, purpose=PURPOSE
    )
    def _invoke(active_llm: LLMProvider):
        return generate_structured(
            active_llm,
            schema=_LLMDraftResponse,
            messages=messages,
            context=llm_context,
            max_tokens=max_tokens_for_purpose(PURPOSE_DRAFTING),
            session=session,
        )

    try:
        response, completion = _invoke(llm)
    except LLMQuotaExhaustedError as quota_exc:
        # Hard cutover: Anthropic returned 402 ("credit balance is too
        # low"). Retrying on Haiku would hit the same wall, so we go
        # straight to the OpenAI cross-provider fallback (gpt-5.1).
        logger.warning(
            "Draft primary %s quota exhausted; cutting over to OpenAI: %s",
            getattr(llm, "model", "<unknown>"),
            quota_exc,
        )
        response, completion = _generate_draft_via_openai(_invoke, quota_exc)
    except (LLMProviderError, ValidationError) as exc:
        # Broadened from LLMResponseFormatError (Hari-III-BUG-019 +
        # Ram-BUG-007, 2026-04-22): Anthropic 503s / httpx timeouts / connection
        # errors are wrapped in LLMProviderError (the parent), NOT the
        # format-error child. Catching only the child let 503s escape
        # past the Haiku retry and surface as opaque 500s — users saw
        # "Could not generate a new version." with no actionable
        # detail. Catching the parent means every recoverable upstream
        # failure (overload, malformed JSON, schema mismatch) triggers
        # the Haiku fallback, and only a Haiku-also-down scenario
        # reaches the OpenAI cutover (and only after that the 422).
        logger.warning(
            "Draft LLM %s refused / malformed / upstream error: %s",
            getattr(llm, "model", "<unknown>"),
            exc,
        )
        fallback = _haiku_fallback_provider()
        if fallback is None:
            response, completion = _generate_draft_via_openai(_invoke, exc)
        else:
            try:
                response, completion = _invoke(fallback)
            except LLMQuotaExhaustedError as quota_exc:
                logger.warning(
                    "Draft Haiku fallback hit quota wall; cutting over to OpenAI: %s",
                    quota_exc,
                )
                response, completion = _generate_draft_via_openai(_invoke, quota_exc)
            except (LLMProviderError, ValidationError) as retry_exc:
                logger.warning("Draft Haiku fallback also failed: %s", retry_exc)
                response, completion = _generate_draft_via_openai(_invoke, retry_exc)

    surviving, verified_count = _verify_version_citations(session, response.citations)
    if verified_count == 0 and response.citations:
        logger.info(
            "Draft %s generated with %d citations, 0 verified.",
            draft.id,
            len(response.citations),
        )

    findings = run_validators(response.body, surviving)
    for f in findings:
        logger.warning(
            "Draft validation finding [%s:%s] on draft %s: %s",
            f.severity, f.code, draft.id, f.message,
        )
    augmented_summary = _augment_summary_with_findings(response.summary, findings)

    model_run = _write_model_run(
        session,
        context=context,
        matter_id=matter.id,
        completion=completion,
        prompt_hash=prompt_hash,
    )

    next_revision = 1
    if draft.versions:
        next_revision = max(v.revision for v in draft.versions) + 1

    version = DraftVersion(
        draft_id=draft.id,
        generated_by_membership_id=context.membership.id,
        model_run_id=model_run.id,
        revision=next_revision,
        body=response.body,
        citations_json=json.dumps(surviving),
        verified_citation_count=verified_count,
        summary=augmented_summary,
    )
    session.add(version)
    session.flush()

    draft.current_version_id = version.id
    draft.status = DraftStatus.DRAFT  # fresh version always resets to draft
    draft.review_required = True
    draft.updated_at = datetime.now(UTC)
    session.flush()
    record_from_context(
        session,
        context,
        action="draft.version_generated",
        target_type="draft",
        target_id=draft.id,
        matter_id=matter.id,
        metadata={
            "revision": version.revision,
            "verified_citation_count": version.verified_citation_count,
            "cited_count": len(surviving),
        },
    )
    session.commit()
    session.refresh(draft)
    return draft


def _record_review(
    session: Session,
    *,
    draft: Draft,
    version_id: str | None,
    action: str,
    context: SessionContext,
    notes: str | None,
) -> DraftReview:
    review = DraftReview(
        draft_id=draft.id,
        version_id=version_id,
        actor_membership_id=context.membership.id,
        action=action,
        notes=notes,
    )
    session.add(review)
    session.flush()
    return review


def _assert_current_version(draft: Draft) -> DraftVersion:
    if draft.current_version_id is None or not draft.versions:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Draft has no generated version yet. Generate one first.",
        )
    current = next(
        (v for v in draft.versions if v.id == draft.current_version_id), None
    )
    if current is None:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Draft's current version is missing.",
        )
    return current


def transition_draft(
    session: Session,
    *,
    context: SessionContext,
    matter_id: str,
    draft_id: str,
    action: Literal["submit", "request_changes", "approve", "finalize"],
    notes: str | None = None,
) -> Draft:
    matter = _load_matter(session, context, matter_id)
    draft = _load_draft(session, matter, draft_id)

    if draft.status == DraftStatus.FINALIZED:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Draft is finalized; no further transitions allowed.",
        )

    current = _assert_current_version(draft)

    if action == DraftReviewAction.SUBMIT:
        if draft.status not in {DraftStatus.DRAFT, DraftStatus.CHANGES_REQUESTED}:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=f"Cannot submit from status {draft.status!r}.",
            )
        draft.status = DraftStatus.IN_REVIEW
        draft.review_required = True
    elif action == DraftReviewAction.REQUEST_CHANGES:
        if draft.status != DraftStatus.IN_REVIEW:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="Only in-review drafts can have changes requested.",
            )
        draft.status = DraftStatus.CHANGES_REQUESTED
        draft.review_required = True
    elif action == DraftReviewAction.APPROVE:
        if draft.status != DraftStatus.IN_REVIEW:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="Only in-review drafts can be approved.",
            )
        if current.verified_citation_count <= 0:
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=(
                    "Cannot approve a draft with zero verified citations. "
                    "Regenerate with grounded authorities first."
                ),
            )
        draft.status = DraftStatus.APPROVED
        draft.review_required = False
    elif action == DraftReviewAction.FINALIZE:
        if draft.status != DraftStatus.APPROVED:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail="Only approved drafts can be finalized.",
            )
        draft.status = DraftStatus.FINALIZED
        draft.review_required = False
    else:  # pragma: no cover — guarded by literal type
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unknown action {action!r}.",
        )

    _record_review(
        session,
        draft=draft,
        version_id=current.id,
        action=action,
        context=context,
        notes=notes,
    )
    draft.updated_at = datetime.now(UTC)
    session.flush()
    record_from_context(
        session,
        context,
        action=f"draft.{action}",
        target_type="draft",
        target_id=draft.id,
        matter_id=matter.id,
        metadata={
            "status_after": draft.status,
            "version_id": current.id,
            "notes_len": len(notes) if notes else 0,
        },
    )
    session.commit()
    session.refresh(draft)
    return draft


def render_version_docx(
    session: Session,
    *,
    context: SessionContext,
    matter_id: str,
    draft_id: str,
    version_id: str | None = None,
) -> tuple[bytes, str]:
    """Return (docx_bytes, suggested_filename). Falls back to the
    draft's current version when version_id is not supplied."""
    matter = _load_matter(session, context, matter_id)
    draft = _load_draft(session, matter, draft_id)
    target_id = version_id or draft.current_version_id
    if not target_id or not draft.versions:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Draft has no version to export. Generate one first.",
        )
    version = next((v for v in draft.versions if v.id == target_id), None)
    if version is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Draft version not found.",
        )

    # PRD §6.1 / §17.4: a legal draft must be citation-grounded or
    # refused. The `approve` transition already enforces this, but the
    # DOCX export path used to be reachable without any verified
    # citation — a reviewer could download a zero-citation brief and
    # circulate it. Close the loop: export is gated unless at least one
    # citation verified OR the draft has reached approved/finalized
    # (the reviewer has accepted the gap on record).
    gate_bypassed = draft.status in {DraftStatus.APPROVED, DraftStatus.FINALIZED}
    if not gate_bypassed and (version.verified_citation_count or 0) <= 0:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "This draft version has zero verified citations. Export is "
                "refused until at least one citation is verified OR the "
                "reviewing partner explicitly approves the draft on record."
            ),
        )

    # Import inline to avoid pulling python-docx into the process when
    # the export route is never hit (keeps cold start slim).
    from docx import Document  # type: ignore
    from docx.shared import Pt

    doc = Document()
    title = doc.add_heading(draft.title, level=1)
    for run in title.runs:
        run.font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.add_run(f"Matter: {matter.title} ({matter.matter_code}) · ")
    meta.add_run(f"Draft type: {draft.draft_type} · ")
    meta.add_run(f"Revision {version.revision} · ")
    meta.add_run(f"Status: {draft.status}")
    meta.runs[-1].italic = True

    doc.add_paragraph()  # spacer

    # Body — split on blank lines to produce readable paragraphs.
    for block in version.body.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        para = doc.add_paragraph()
        for line_idx, line in enumerate(block.split("\n")):
            if line_idx:
                para.add_run().add_break()
            para.add_run(line)

    try:
        citations = json.loads(version.citations_json) if version.citations_json else []
    except json.JSONDecodeError:
        citations = []
    if citations:
        doc.add_heading("Authorities cited", level=2)
        for c in citations:
            doc.add_paragraph(c, style="List Bullet")

    if draft.review_required:
        doc.add_paragraph(
            "Review required — this draft has not been approved by a partner.",
        ).runs[-1].italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    safe_title = "".join(
        ch if ch.isalnum() or ch in ("-", "_") else "-" for ch in draft.title
    ).strip("-")[:60] or "draft"
    filename = f"{safe_title}-r{version.revision}.docx"
    return buffer.getvalue(), filename


def load_draft_record(draft: Draft) -> dict:
    """Shape the Draft for pydantic serialisation with parsed citations."""
    versions = sorted(draft.versions, key=lambda v: v.revision)
    versions_payload = []
    for v in versions:
        try:
            cites = json.loads(v.citations_json) if v.citations_json else []
        except json.JSONDecodeError:
            cites = []
        versions_payload.append(
            {
                "id": v.id,
                "draft_id": v.draft_id,
                "revision": v.revision,
                "body": v.body,
                "citations": cites,
                "verified_citation_count": v.verified_citation_count,
                "summary": v.summary,
                "generated_by_membership_id": v.generated_by_membership_id,
                "model_run_id": v.model_run_id,
                "created_at": v.created_at,
            }
        )
    return {
        "id": draft.id,
        "matter_id": draft.matter_id,
        "created_by_membership_id": draft.created_by_membership_id,
        "title": draft.title,
        "draft_type": draft.draft_type,
        "template_type": draft.template_type,
        "status": draft.status,
        "review_required": draft.review_required,
        "current_version_id": draft.current_version_id,
        "versions": versions_payload,
        "reviews": [
            {
                "id": r.id,
                "draft_id": r.draft_id,
                "version_id": r.version_id,
                "actor_membership_id": r.actor_membership_id,
                "action": r.action,
                "notes": r.notes,
                "created_at": r.created_at,
            }
            for r in sorted(draft.reviews, key=lambda r: r.created_at)
        ],
        "created_at": draft.created_at,
        "updated_at": draft.updated_at,
    }
